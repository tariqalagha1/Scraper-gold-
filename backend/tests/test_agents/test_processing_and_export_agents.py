from pathlib import Path
from uuid import uuid4

import pytest

try:
    from docx import Document as DocxDocument
except ModuleNotFoundError:  # pragma: no cover - optional dependency in CI
    DocxDocument = None

from app.agents.analysis_agent import AnalysisAgent
from app.agents.export_agent import ExportAgent
from app.agents.processing_agent import ProcessingAgent
from app.agents.vector_agent import VectorAgent
from app.config import settings
from app.export.word_exporter import WordExporter
from app.scraper.extractor import ContentExtractor
from app.storage.manager import StorageManager

pytestmark = pytest.mark.asyncio


SAMPLE_HTML = """
<html>
  <head><title>Agent Test Page</title></head>
  <body>
    <h1>Agent Test Page</h1>
    <p>This processing fixture contains enough text to validate normalization and exports.</p>
    <a href="https://example.com/report.pdf">Report PDF</a>
  </body>
</html>
"""

LIST_HTML = """
<html>
  <body>
    <section>
      <article class="product_pod">
        <h3><a href="/catalogue/its-only-the-himalayas_981/index.html">It's Only the Himalayas</a></h3>
        <p class="price_color">GBP45.17</p>
      </article>
      <article class="product_pod">
        <h3><a href="/catalogue/full-moon-over-noahs-ark-an-odyssey-to-mount-ararat-and-beyond_811/index.html">Full Moon over Noah's Ark</a></h3>
        <p class="price_color">GBP49.43</p>
      </article>
      <article class="product_pod">
        <h3><a href="/catalogue/see-america-a-celebration-of-our-national-parks-treasured-sites-and-festivals_732/index.html">See America</a></h3>
        <p class="price_color">GBP48.87</p>
      </article>
      <article class="product_pod">
        <h3><a href="/catalogue/vagabonding-an-uncommon-guide-to-the-art-of-long-term-world-travel_615/index.html">Vagabonding</a></h3>
        <p class="price_color">GBP36.94</p>
      </article>
      <article class="product_pod">
        <h3><a href="/catalogue/under-the-tuscan-sun_504/index.html">Under the Tuscan Sun</a></h3>
        <p class="price_color">GBP37.33</p>
      </article>
    </section>
  </body>
</html>
"""


async def test_processing_agent_returns_normalized_payload_with_metadata():
    extractor = ContentExtractor()
    extracted = extractor.extract(raw_html=SAMPLE_HTML, url="https://example.com/test")

    result = await ProcessingAgent().safe_execute({"extracted": extracted, "url": "https://example.com/test"})

    assert result["status"] == "success"
    assert result["data"]["cleaned_text"]
    assert result["data"]["source_url"] == "https://example.com/test"
    assert result["data"]["items"][0]["content"]
    assert result["metadata"]["agent"] == "processing_agent"
    assert result["metadata"]["source"] == "https://example.com/test"
    assert result["metadata"]["type"] == "processing_agent"
    assert result["metadata"]["execution_time"]


async def test_processing_agent_returns_structured_list_items_for_books_listing():
    extractor = ContentExtractor()
    extracted = extractor.extract(
        raw_html=LIST_HTML,
        url="https://books.toscrape.com/catalogue/category/books/travel_2/index.html",
        selectors={
            "container": "article.product_pod",
            "fields": {
                "title": "h3 a",
                "link": "h3 a[href]",
                "price": ".price_color",
            },
        },
    )

    result = await ProcessingAgent().safe_execute(
        {
            "extracted": extracted,
            "url": "https://books.toscrape.com/catalogue/category/books/travel_2/index.html",
        }
    )

    assert result["status"] == "success"
    assert result["data"]["page_type"] == "list_page"
    assert len(result["data"]["items"]) == 5
    assert result["data"]["items"][0]["title"] == "It's Only the Himalayas"
    assert result["data"]["items"][0]["price"] == "GBP45.17"
    assert result["data"]["items"][0]["link"].endswith("/catalogue/its-only-the-himalayas_981/index.html")


async def test_processing_agent_preserves_ai_selected_record_fields():
    html = """
    <html>
      <body>
        <article class="product-card">
          <h2><a href="/item-1">Item One</a></h2>
          <span class="price">$10</span>
          <span class="availability">In stock</span>
        </article>
      </body>
    </html>
    """
    extractor = ContentExtractor()
    extracted = extractor.extract(
        raw_html=html,
        url="https://example.com/catalog",
        selectors={
            "container": "article.product-card",
            "fields": {
                "title": "h2 a",
                "link": "h2 a[href]",
                "price": ".price",
                "availability": ".availability",
            },
        },
    )

    result = await ProcessingAgent().safe_execute(
        {
            "extracted": extracted,
            "url": "https://example.com/catalog",
        }
    )

    assert result["status"] == "success"
    assert result["data"]["items"][0]["availability"] == "In stock"
    assert "In stock" in result["data"]["items"][0]["content"]


async def test_processing_agent_honors_max_records_limit():
    extractor = ContentExtractor()
    extracted = extractor.extract(
        raw_html=LIST_HTML,
        url="https://books.toscrape.com/catalogue/category/books/travel_2/index.html",
        selectors={
            "container": "article.product_pod",
            "fields": {
                "title": "h3 a",
                "link": "h3 a[href]",
                "price": ".price_color",
            },
        },
    )

    result = await ProcessingAgent().safe_execute(
        {
            "extracted": extracted,
            "url": "https://books.toscrape.com/catalogue/category/books/travel_2/index.html",
            "max_records": 2,
        }
    )

    assert result["status"] == "success"
    assert len(result["data"]["items"]) == 2


async def test_processing_agent_projects_requested_record_fields_only():
    html = """
    <html>
      <body>
        <table>
          <thead>
            <tr><th>Registration No</th><th>Name</th><th>Phone Number</th><th>Status</th></tr>
          </thead>
          <tbody>
            <tr><td>1001</td><td>Fatimah Ali</td><td>0500000001</td><td>Active</td></tr>
            <tr><td>1002</td><td>Mona Saeed</td><td>0500000002</td><td>Active</td></tr>
            <tr><td>1003</td><td>Layan Omar</td><td>0500000003</td><td>Inactive</td></tr>
          </tbody>
        </table>
      </body>
    </html>
    """

    extractor = ContentExtractor()
    extracted = extractor.extract(raw_html=html, url="https://example.com/patients")
    result = await ProcessingAgent().safe_execute(
        {
            "extracted": extracted,
            "url": "https://example.com/patients",
            "record_fields": ["name"],
        }
    )

    assert result["status"] == "success"
    assert len(result["data"]["items"]) == 3
    assert sorted(result["data"]["items"][0].keys()) == ["content", "name", "source_url", "type"]
    assert result["data"]["items"][0]["name"] == "Fatimah Ali"


async def test_export_agent_generates_all_formats_from_pipeline_payload(isolated_storage):
    extractor = ContentExtractor()
    processed = await ProcessingAgent().safe_execute(
        {"extracted": extractor.extract(raw_html=SAMPLE_HTML, url="https://example.com/test"), "url": "https://example.com/test"}
    )
    analysis = await AnalysisAgent().safe_execute({"items": processed["data"]["items"], "context": "fixture"})

    export_id = f"export-{uuid4()}"
    result = await ExportAgent().safe_execute(
        {
            "processed": processed,
            "analysis": analysis["data"],
            "source_url": "https://example.com/test",
            "export_id": export_id,
            "title": "Pipeline Export Verification",
        }
    )

    assert result["status"] == "success"
    assert result["data"]["analysis_included"] is True
    assert result["metadata"]["source"] == "https://example.com/test"
    assert result["metadata"]["execution_time"]

    storage = StorageManager()
    expected_extensions = {
        "excel_path": ".xlsx",
        "pdf_path": ".pdf",
        "word_path": ".docx",
    }
    for key, extension in expected_extensions.items():
        export_path = result["data"][key]
        assert export_path
        assert storage.file_exists(export_path)
        assert storage.get_file_size(export_path) > 0
        assert Path(export_path).suffix.lower() == extension
        assert Path(export_path).parts[0] == "exports"


async def test_word_exporter_includes_full_record_set_and_clean_empty_errors(isolated_storage):
    if DocxDocument is None:
        pytest.skip("python-docx is not installed")

    records = [
        {"name": f"Patient {index}", "id": str(1000 + index), "mobile": f"050000{index:04d}"}
        for index in range(12)
    ]
    payload = {
        "status": "completed",
        "request": {"url": "https://example.com/patients"},
        "result": {
            "data": records,
            "raw": {},
            "processed": {"summary": "ok"},
            "analysis": {},
            "vector": {},
            "exports": {},
        },
        "execution": {
            "decision": {"page_type": "list", "confidence": 0.9, "reason": "fixture"},
            "validation": {"status": "pass", "confidence": 0.9, "issues": [], "metrics": {}, "should_retry": False},
            "retry": {"attempted": False, "result": False},
            "memory": {"used": False, "selector_source": "generated", "success_rate": None},
            "timing": {},
            "steps": {"current": "export"},
            "trace": {},
        },
        "metadata": {"run_id": "run-export-test", "job_id": "job-export-test", "user_id": "user-export-test"},
        "errors": [],
    }

    exporter = WordExporter()
    exported_path = await exporter.export(payload, source_url="https://example.com/patients", title="Patients")

    storage = StorageManager()
    absolute_path = storage.resolve_path(exported_path)
    document = DocxDocument(str(absolute_path))

    assert document.tables
    assert len(document.tables[0].rows) == len(records) + 1
    assert any(paragraph.text.strip() == "No errors reported." for paragraph in document.paragraphs)


async def test_vector_agent_skips_when_disabled(monkeypatch):
    monkeypatch.setattr(settings, "ENABLE_VECTOR", False)
    monkeypatch.setattr(settings, "OPENAI_API_KEY", "sk-test")

    async def fail_generate(*args, **kwargs):
        raise AssertionError("Embedding provider should not be called when vector is disabled")

    agent = VectorAgent()
    monkeypatch.setattr(agent.embedding_generator, "generate", fail_generate)

    result = await agent.safe_execute(
        {
            "operation": "embed",
            "items": [{"content": "Example content", "source_url": "https://example.com/test"}],
        }
    )

    assert result["status"] == "success"
    assert result["data"]["status"] == "skipped"
    assert result["data"]["reason"] == "vector_store_disabled"


async def test_vector_agent_skips_when_openai_api_key_missing(monkeypatch):
    monkeypatch.setattr(settings, "ENABLE_VECTOR", True)
    monkeypatch.setattr(settings, "OPENAI_API_KEY", "")

    async def fail_generate(*args, **kwargs):
        raise AssertionError("Embedding provider should not be called when API key is missing")

    agent = VectorAgent()
    monkeypatch.setattr(agent.embedding_generator, "generate", fail_generate)

    result = await agent.safe_execute(
        {
            "operation": "embed",
            "items": [{"content": "Example content", "source_url": "https://example.com/test"}],
        }
    )

    assert result["status"] == "success"
    assert result["data"]["status"] == "skipped"
    assert result["data"]["reason"] == "embedding_provider_unconfigured"
