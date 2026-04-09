from __future__ import annotations

from io import BytesIO

from docx import Document

from app.export.base_exporter import BaseExporter
from app.export.contract_helpers import (
    get_export_data,
    get_export_errors,
    get_export_execution_summary,
    get_export_metadata,
    normalize_export_contract,
)


class WordExporter(BaseExporter):
    async def export(
        self,
        processed_data: dict,
        *,
        analysis_data: dict | None = None,
        export_id: str | None = None,
        source_url: str = "",
        generated_at: str | None = None,
        title: str = "",
    ) -> str:
        resolved_export_id = self.resolve_export_id(export_id)
        resolved_generated_at = self.resolve_generated_at(generated_at)
        contract = normalize_export_contract(processed_data, analysis_data=analysis_data, source_url=source_url)
        metadata = get_export_metadata(contract)
        execution = get_export_execution_summary(contract)
        records = get_export_data(contract)
        errors = get_export_errors(contract)

        document = Document()
        document.add_heading(title or "Extraction Report", 0)
        document.add_paragraph(f"Source: {metadata['URL']}")
        document.add_paragraph(f"Status: {contract.get('status', '')}")
        document.add_paragraph(f"Run ID: {metadata['Run ID']}")
        document.add_paragraph(f"Date: {resolved_generated_at}")

        document.add_heading("Execution Summary", level=1)
        document.add_paragraph(f"Items extracted: {len(records)}")
        document.add_paragraph(f"Validation status: {execution['Validation Status']}")
        document.add_paragraph(f"Validation confidence: {execution['Validation Confidence']}")
        document.add_paragraph(f"Retry attempted: {execution['Retry Attempted']}")
        document.add_paragraph(f"Retry result: {execution['Retry Result']}")
        document.add_paragraph(f"Memory used: {execution['Memory Used']}")

        document.add_heading("Decision", level=1)
        document.add_paragraph(f"Page Type: {execution['Page Type']}")
        document.add_paragraph(f"Decision Confidence: {execution['Decision Confidence']}")
        document.add_paragraph(f"Reason: {execution['Decision Reason'] or 'None'}")

        document.add_heading("Data Preview", level=1)
        preview_rows = records[:10]
        if preview_rows:
            headers = sorted({key for row in preview_rows for key in row.keys()})[:5]
            table = document.add_table(rows=1, cols=max(len(headers), 1))
            table.style = "Table Grid"
            header_cells = table.rows[0].cells
            for idx, value in enumerate(headers or ["Value"]):
                header_cells[idx].text = str(value)
            for row_data in preview_rows:
                row_cells = table.add_row().cells
                for idx, header in enumerate(headers[: len(row_cells)]):
                    row_cells[idx].text = str(row_data.get(header, ""))
            document.add_paragraph("")
        else:
            document.add_paragraph("No structured data available.")

        document.add_heading("Errors", level=1)
        for error in errors:
            document.add_paragraph(error, style="List Bullet")

        buffer = BytesIO()
        document.save(buffer)
        return self.write_export_file("docx", buffer.getvalue(), export_id=resolved_export_id)
